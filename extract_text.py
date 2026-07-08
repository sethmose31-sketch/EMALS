"""
extract_text.py
----------------
Pulls raw text out of CV files (PDF and DOCX). This is intentionally simple
and format-agnostic — all the "intelligence" (structuring, correcting typos,
inferring missing fields) happens later in gemini_processor.py.
"""

import os
import pdfplumber
from docx import Document


def extract_from_pdf(path: str) -> str:
    """Extract all text from a PDF file, page by page."""
    text_chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_from_docx(path: str) -> str:
    """Extract all text from a DOCX file, including tables (some CVs use
    tables for layout, e.g. a two-column skills section)."""
    doc = Document(path)
    text_chunks = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_chunks.append(row_text)

    return "\n".join(text_chunks)


def extract_text(path: str) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_from_pdf(path)
    elif ext in (".docx", ".doc"):
        if ext == ".doc":
            raise ValueError(
                f"'.doc' (legacy Word format) is not supported: {path}. "
                "Please convert it to .docx or .pdf first."
            )
        return extract_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {path}")


def collect_cv_files(folder: str) -> list:
    """Return a sorted list of all supported CV file paths in a folder."""
    supported = (".pdf", ".docx")
    files = [
        os.path.join(folder, f)
        for f in os.listdir(folder)
        if f.lower().endswith(supported) and not f.startswith("~$")
    ]
    return sorted(files)
